import os
import hashlib
import asyncio
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import fitz  # PyMuPDF
import PyPDF2
from docx import Document as DocxDocument
import chardet
from dataclasses import dataclass
from loguru import logger
from .config import get_settings

settings = get_settings()


@dataclass
class DocumentChunk:
    content: str
    page_number: Optional[int] = None
    start_char: Optional[int] = None
    end_char: Optional[int] = None
    metadata: Optional[Dict[str, Any]] = None


@dataclass
class ProcessedDocument:
    filename: str
    file_type: str
    file_size: int
    content_hash: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_time: float


class DocumentProcessor:
    """Advanced document processing with multiple extraction methods"""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.supported_types = settings.allowed_file_types
    
    async def process_file(self, file_path: str, filename: str) -> ProcessedDocument:
        """Process uploaded file and extract text content"""
        start_time = asyncio.get_event_loop().time()
        
        file_type = self._get_file_type(filename)
        file_size = os.path.getsize(file_path)
        
        # Extract text content
        if file_type == "pdf":
            content, metadata = await self._extract_pdf_content(file_path)
        elif file_type == "docx":
            content, metadata = await self._extract_docx_content(file_path)
        elif file_type == "txt":
            content, metadata = await self._extract_text_content(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Generate content hash
        content_hash = self._generate_content_hash(content)
        
        # Create chunks
        chunks = self._create_chunks(content, metadata)
        
        processing_time = asyncio.get_event_loop().time() - start_time
        
        return ProcessedDocument(
            filename=filename,
            file_type=file_type,
            file_size=file_size,
            content_hash=content_hash,
            chunks=chunks,
            metadata=metadata,
            processing_time=processing_time
        )
    
    def _get_file_type(self, filename: str) -> str:
        """Determine file type from extension"""
        ext = Path(filename).suffix.lower().lstrip('.')
        if ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {ext}")
        return ext
    
    async def _extract_pdf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from PDF using multiple methods for robustness"""
        content_parts = []
        metadata = {"pages": [], "total_pages": 0}
        
        try:
            # Primary method: PyMuPDF (better for complex layouts)
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    content_parts.append(text)
                    metadata["pages"].append({
                        "page_number": page_num + 1,
                        "char_count": len(text)
                    })
            
            metadata["total_pages"] = len(doc)
            metadata["extraction_method"] = "PyMuPDF"
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}, trying PyPDF2")
            
            # Fallback method: PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            content_parts.append(text)
                            metadata["pages"].append({
                                "page_number": page_num + 1,
                                "char_count": len(text)
                            })
                    
                    metadata["total_pages"] = len(reader.pages)
                    metadata["extraction_method"] = "PyPDF2"
                    
            except Exception as e2:
                raise Exception(f"PDF extraction failed with both methods: {e}, {e2}")
        
        content = "\n\n".join(content_parts)
        return content, metadata
    
    async def _extract_docx_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from DOCX files"""
        try:
            doc = DocxDocument(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            content = "\n\n".join(content_parts)
            metadata = {
                "paragraph_count": len(content_parts),
                "extraction_method": "python-docx"
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {e}")
    
    async def _extract_text_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from plain text files with encoding detection"""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_info = chardet.detect(raw_data)
                encoding = encoding_info.get('encoding', 'utf-8')
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            metadata = {
                "encoding": encoding,
                "confidence": encoding_info.get('confidence', 0.0),
                "extraction_method": "chardet"
            }
            
            return content, metadata
            
        except Exception as e:
            raise Exception(f"Text extraction failed: {e}")
    
    def _generate_content_hash(self, content: str) -> str:
        """Generate SHA-256 hash of content for deduplication"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()
    
    def _create_chunks(self, content: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Create overlapping text chunks for embedding"""
        if not content.strip():
            return []
        
        chunks = []
        content_length = len(content)
        current_pos = 0
        chunk_index = 0
        
        while current_pos < content_length:
            # Calculate chunk end position
            chunk_end = min(current_pos + self.chunk_size, content_length)
            
            # Adjust chunk boundary to avoid breaking words
            if chunk_end < content_length:
                # Look for sentence endings first
                for punct in ['. ', '! ', '? ', '\n\n']:
                    last_punct = content.rfind(punct, current_pos, chunk_end)
                    if last_punct > current_pos + self.chunk_size // 2:
                        chunk_end = last_punct + len(punct)
                        break
                else:
                    # Fallback to word boundary
                    last_space = content.rfind(' ', current_pos, chunk_end)
                    if last_space > current_pos + self.chunk_size // 2:
                        chunk_end = last_space
            
            # Extract chunk content
            chunk_content = content[current_pos:chunk_end].strip()
            
            if chunk_content:
                # Determine page number from metadata
                page_number = self._get_page_for_position(current_pos, metadata)
                
                chunk = DocumentChunk(
                    content=chunk_content,
                    page_number=page_number,
                    start_char=current_pos,
                    end_char=chunk_end,
                    metadata={"chunk_index": chunk_index}
                )
                
                chunks.append(chunk)
                chunk_index += 1
            
            # Move to next chunk with overlap
            current_pos = chunk_end - self.chunk_overlap
            if current_pos >= chunk_end:  # Prevent infinite loop
                break
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
    
    def _get_page_for_position(self, position: int, metadata: Dict[str, Any]) -> Optional[int]:
        """Estimate page number for a character position"""
        if "pages" not in metadata:
            return None
        
        char_count = 0
        for page_info in metadata["pages"]:
            char_count += page_info.get("char_count", 0)
            if position <= char_count:
                return page_info["page_number"]
        
        return metadata.get("total_pages", 1)
    
    def validate_file(self, file_path: str, filename: str) -> Tuple[bool, str]:
        """Validate uploaded file"""
        try:
            file_type = self._get_file_type(filename)
            file_size = os.path.getsize(file_path)
            
            # Check file size
            max_size_bytes = settings.max_file_size_mb * 1024 * 1024
            if file_size > max_size_bytes:
                return False, f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds limit ({settings.max_file_size_mb}MB)"
            
            # Check if file is readable
            if not os.access(file_path, os.R_OK):
                return False, "File is not readable"
            
            return True, "File is valid"
            
        except ValueError as e:
            return False, str(e)
        except Exception as e:
            return False, f"File validation failed: {e}"